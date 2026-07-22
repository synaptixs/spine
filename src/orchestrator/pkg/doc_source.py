"""Read a repo's documentation files into ``DocPage`` rows for the doc-semantic layer.

The counterpart to the language extractors, for prose: walk a repository for documentation —
markdown, ``.rst``/``.txt``, **HTML**, **PDF**, and **Office** (``.docx``/``.xlsx``) — and hand
each file to ``pkg.docs`` as a ``DocPage`` (title = repo-relative path, so mentions resolve
relative to the doc). Deterministic, no LLM, no network.

Formats **register** rather than branch: each is a :class:`DocReader` on the module registry
(see :func:`register_reader`), so adding one touches no existing reader and no dispatch code.
A reader returning ``None`` means *skip this file* — too large, unparseable, or an optional
dependency is missing. Skipping is always non-fatal: an unreadable format is absent from the
graph, never an error.

Markdown, text, and HTML parse with the stdlib. **PDF** needs ``pypdf`` behind the optional
``[docs]`` extra; **Office** needs ``python-docx``/``openpyxl`` behind ``[office]``. Both are
lazy-imported so the base install stays stdlib-only, and an absent extra just means those files
are skipped — as is a scanned PDF (no extractable text; we don't OCR) or an encrypted document.
"""

from __future__ import annotations

import os
import re
from collections.abc import Callable
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path

from orchestrator.pkg.docs import DocPage
from orchestrator.pkg.extractor import DEFAULT_IGNORE_DIRS

# Skip absurdly large docs (generated dumps, vendored changelogs) — keep the graph legible.
_MAX_DOC_BYTES = 1_000_000
# PDFs are binary and legitimately larger; cap bytes and pages so a giant scan can't stall a walk.
_MAX_PDF_BYTES = 25_000_000
_MAX_PDF_PAGES = 500
# Cap section-granular nodes per doc: a runaway doc (hundreds of headings) shouldn't flood the
# graph. Beyond this, the doc stays whole rather than exploding into fragments.
_MAX_SECTIONS = 40
# An ATX markdown heading: `#`..`######` then text. Setext / RST underlines aren't split (they'd
# need lookahead and are rarer in the docs this targets); those docs stay whole — safe, not wrong.
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")


@dataclass(frozen=True)
class DocReader:
    """How one family of file suffixes becomes a doc's text.

    ``read`` returns the doc's text, or ``None`` to **skip** the file — too large, unparseable,
    or an optional dependency isn't installed. Skipping is always non-fatal: a format we can't
    read is simply absent from the graph, never an error (the ``pypdf`` precedent).

    ``sections=True`` marks formats whose text carries markdown-style ATX headings, so
    :func:`split_sections` can split them into section-granular pages. A format that doesn't
    (plain text, PDF) stays one page per file.
    """

    name: str
    suffixes: frozenset[str]
    read: Callable[[Path], str | None]
    sections: bool = False


_READERS: dict[str, DocReader] = {}


def register_reader(reader: DocReader) -> None:
    """Register ``reader`` for each of its suffixes (last registration wins).

    This is the seam new formats plug into: adding one touches no existing reader and no
    dispatch branch. Suffixes are matched lowercased.
    """
    for suffix in reader.suffixes:
        _READERS[suffix.lower()] = reader


def is_doc_file(path: Path) -> bool:
    """True when some registered reader claims this suffix.

    A claimed suffix still may not *read* — PDF needs the ``[docs]`` extra — but it is a
    documentation file as far as the walker is concerned.
    """
    return path.suffix.lower() in _READERS


def _slug(heading: str) -> str:
    """A GitHub-style anchor slug for a heading (lowercase, spaces→dashes, punctuation dropped)."""
    text = re.sub(r"[^\w\s-]", "", heading.strip().lower())
    return re.sub(r"[\s_]+", "-", text).strip("-")


def split_sections(page: DocPage) -> list[DocPage]:
    """Split a markdown ``DocPage`` into one page per heading (``path#slug``), or ``[page]`` unchanged.

    Section granularity lets a ``MENTIONS`` edge point at the *section* that names a symbol, not the
    whole file, and gives each section provenance at its heading line. Bounded by ``_MAX_SECTIONS``:
    a doc with more headings than that stays whole rather than fragmenting the graph. Content before
    the first heading (a preamble) becomes a section keyed by the bare path, so nothing is dropped.
    Deterministic; unique slugs are disambiguated with a numeric suffix."""
    lines = page.text.splitlines()
    heads = [(i, m) for i, line in enumerate(lines) if (m := _HEADING_RE.match(line))]
    if not heads or len(heads) > _MAX_SECTIONS:
        return [page]

    file = page.source_file or page.title
    bounds = [i for i, _ in heads] + [len(lines)]
    sections: list[DocPage] = []
    seen: dict[str, int] = {}
    # Preamble before the first heading (if any real content) → a page keyed by the bare path.
    if heads[0][0] > 0 and "".join(lines[: heads[0][0]]).strip():
        pre = "\n".join(lines[: heads[0][0]])
        sections.append(DocPage(title=file, text=pre, base_dir=page.base_dir, source_file=file, line=1))
    for idx, (line_no, m) in enumerate(heads):
        slug = _slug(m.group(2)) or "section"
        seen[slug] = seen.get(slug, 0) + 1
        if seen[slug] > 1:
            slug = f"{slug}-{seen[slug]}"
        body = "\n".join(lines[line_no : bounds[idx + 1]])
        title = f"{file}#{slug}"
        sections.append(
            DocPage(title=title, text=body, base_dir=page.base_dir, source_file=file, line=line_no + 1)
        )
    return sections


def _read_pdf_text(path: Path) -> str | None:
    """A PDF's text (pages joined by blank lines), or ``None`` if it can't be read.

    ``None`` covers three cases, all non-fatal: ``pypdf`` isn't installed (the ``[docs]``
    extra), the file isn't a parseable PDF, or it has no extractable text (a scanned image —
    we don't OCR). pypdf's error surface on malformed input is wide, so the parse is
    best-effort per page."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return None
    try:
        reader = PdfReader(str(path))
        parts: list[str] = []
        for page in reader.pages[:_MAX_PDF_PAGES]:
            parts.append(page.extract_text() or "")
    except Exception:  # noqa: BLE001 — malformed/encrypted PDFs raise assorted pypdf errors
        return None
    text = "\n\n".join(p for p in parts if p.strip())
    return text or None


def read_doc_pages(root: Path | str, *, sections: bool = True) -> list[DocPage]:
    """Every documentation file under ``root`` as a ``DocPage`` (repo-relative title).

    Skips the usual ignored dirs (``.git``, ``node_modules``, build output, hidden dirs),
    text files that can't be read as UTF-8 or exceed ``_MAX_DOC_BYTES``, and PDFs that are
    too large or yield no text. Deterministic order.

    With ``sections=True`` (the default), markdown docs are split by heading into section-granular
    pages (``path#slug``) via :func:`split_sections`; other formats and heading-less docs stay
    whole. Pass ``sections=False`` for strictly one page per file."""
    root_path = Path(root).resolve()
    pages: list[DocPage] = []
    for dirpath, dirnames, filenames in os.walk(root_path):
        dirnames[:] = sorted(d for d in dirnames if d not in DEFAULT_IGNORE_DIRS and not d.startswith("."))
        for name in sorted(filenames):
            path = Path(dirpath) / name
            reader = _READERS.get(path.suffix.lower())
            if reader is None:
                continue
            text = reader.read(path)
            if text is None:
                continue
            rel = path.relative_to(root_path).as_posix()
            page = DocPage(title=rel, text=text, base_dir=Path(rel).parent.as_posix(), source_file=rel)
            if sections and reader.sections:
                pages.extend(split_sections(page))
            else:
                pages.append(page)
    return pages


def _read_text(path: Path) -> str | None:
    try:
        if path.stat().st_size > _MAX_DOC_BYTES:
            return None
        return path.read_text(encoding="utf-8")
    except (OSError, UnicodeDecodeError):
        return None


def _read_pdf(path: Path) -> str | None:
    try:
        if path.stat().st_size > _MAX_PDF_BYTES:
            return None
    except OSError:
        return None
    return _read_pdf_text(path)


# ---- HTML -------------------------------------------------------------------

_HTML_SKIP_TAGS = frozenset({"script", "style", "noscript", "template", "svg"})
# Tags after which a line break keeps the flattened text readable (and keeps a heading on
# its own line, which is what makes the markdown-heading trick below work).
_HTML_BLOCK_TAGS = frozenset(
    {
        "p", "div", "section", "article", "header", "footer", "aside", "main", "nav",
        "ul", "ol", "li", "table", "tr", "td", "th", "br", "hr", "pre", "blockquote",
        "dl", "dt", "dd", "figcaption",
    }
)  # fmt: skip
_HTML_HEADING_TAG_RE = re.compile(r"h([1-6])")
# Inline code tags → backticks. `<code>Invoice</code>` is HTML's way of saying `Invoice`, and
# the binder treats a backticked mention as a high-confidence code claim. Flattening these to
# bare words would throw away the strongest signal an HTML doc carries.
_HTML_CODE_TAGS = frozenset({"code", "tt", "kbd", "samp", "var"})
_CODE_SPAN_RE = re.compile(r"`\s*([^`\n]*?)\s*`")


class _HtmlToText(HTMLParser):
    """Flatten HTML to markdown-ish text, turning ``<h1>``…``<h6>`` into ATX headings.

    Emitting ``#``-style headings is the whole trick: ``split_sections`` then treats an HTML
    doc exactly like a markdown one, so HTML gets section granularity for free and nothing
    downstream needs to know HTML exists. Inline ``<code>`` becomes a backtick span for the
    same reason — it lands in the binder as the high-confidence mention it already was.

    ``<script>``/``<style>`` content is dropped — it's code, not prose, and would otherwise
    produce a stream of junk mentions.
    """

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._out: list[str] = []
        self._skip = 0
        self._pre = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in _HTML_SKIP_TAGS:
            self._skip += 1
            return
        if tag == "pre":
            self._pre += 1
        if (m := _HTML_HEADING_TAG_RE.fullmatch(tag)) is not None:
            self._out.append("\n" + "#" * int(m.group(1)) + " ")
        elif tag in _HTML_CODE_TAGS:
            # Only *inline* code becomes a backtick span; inside <pre> it's a whole sample, and
            # wrapping that would produce one enormous backticked "mention" that binds to nothing.
            if not self._pre:
                self._out.append(" `")
        elif tag in _HTML_BLOCK_TAGS:
            self._out.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in _HTML_SKIP_TAGS:
            self._skip = max(0, self._skip - 1)
            return
        if _HTML_HEADING_TAG_RE.fullmatch(tag) is not None:
            self._out.append("\n")
        elif tag in _HTML_CODE_TAGS:
            if not self._pre:
                self._out.append("` ")
        elif tag in _HTML_BLOCK_TAGS:
            self._out.append("\n")
        if tag == "pre":
            self._pre = max(0, self._pre - 1)

    def handle_data(self, data: str) -> None:
        if self._skip:
            return
        text = " ".join(data.split())  # collapse whitespace; block tags supply the newlines
        if text:
            self._out.append(text + " ")

    def text(self) -> str:
        """The flattened text, with code spans tightened and blank-line runs collapsed."""
        # `handle_data` pads each chunk with a trailing space, so a code span arrives as
        # "` Invoice `". Tighten it to "`Invoice`" — the binder matches the span's exact
        # contents, so stray whitespace would stop it resolving. Empty spans are dropped.
        joined = _CODE_SPAN_RE.sub(lambda m: f"`{m.group(1)}`" if m.group(1) else "", "".join(self._out))
        out: list[str] = []
        for raw_line in joined.splitlines():
            line = re.sub(r"[ \t]{2,}", " ", raw_line).strip()
            if line or (out and out[-1]):
                out.append(line)
        return "\n".join(out).strip()


def _read_html(path: Path) -> str | None:
    raw = _read_text(path)  # reuses the size cap + UTF-8 handling
    if raw is None:
        return None
    parser = _HtmlToText()
    try:
        parser.feed(raw)
        parser.close()
    except Exception:  # noqa: BLE001 — malformed markup must not fail the walk
        return None
    return parser.text() or None


# ---- Office (.docx / .xlsx) -------------------------------------------------
# Both need the `[office]` extra, lazy-imported like `pypdf`: absent → the format is simply
# skipped. Office files are ZIP containers, so they're binary — `_read_text` never applies.

_MAX_OFFICE_BYTES = 25_000_000
_MAX_SHEETS = 20
_MAX_SHEET_ROWS = 500
_DOCX_HEADING_RE = re.compile(r"heading\s+([1-6])")
# Fonts Word documents use for inline code. A run in one of these (or in a character style
# named "…Code…") is the .docx equivalent of a markdown backtick.
_DOCX_CODE_FONTS = frozenset(
    {"courier", "courier new", "consolas", "monaco", "menlo", "lucida console", "cascadia mono"}
)


def _docx_heading_level(style_name: str) -> int:
    """Word's ``Heading 1``…``Heading 6`` → 1–6; ``Title`` → 1; anything else → 0.

    Word's built-in heading styles are the structural equivalent of markdown's ``#`` levels,
    so mapping them gives ``.docx`` the same section granularity as markdown and HTML.
    """
    name = style_name.strip().lower()
    if name == "title":
        return 1
    match = _DOCX_HEADING_RE.fullmatch(name)
    return int(match.group(1)) if match else 0


def _docx_paragraph_text(paragraph: object) -> str:
    """A paragraph's text, with code-styled runs wrapped in backticks.

    Word marks inline code either with a character style whose name contains "code" or with a
    monospace font. That is the ``.docx`` equivalent of a markdown backtick or an HTML
    ``<code>``, and the binder treats a backticked mention as a high-confidence claim — so
    preserving it is what makes a real Word spec bind instead of reading as plain prose.

    Runs are concatenated *without* inserting separators: Word splits a single word across runs
    freely (spell-check, formatting), so joining on spaces would fracture identifiers.
    """
    parts: list[str] = []
    for run in getattr(paragraph, "runs", []):
        text: str = run.text or ""
        if not text:
            continue
        style = (getattr(getattr(run, "style", None), "name", "") or "").lower()
        font = (getattr(getattr(run, "font", None), "name", "") or "").lower()
        stripped = text.strip()
        if stripped and ("code" in style or font in _DOCX_CODE_FONTS):
            # replace() keeps the run's own leading/trailing spacing intact.
            parts.append(text.replace(stripped, f"`{stripped}`", 1))
        else:
            parts.append(text)
    raw = "".join(parts) if parts else str(getattr(paragraph, "text", ""))
    return " ".join(raw.split())


def _read_docx(path: Path) -> str | None:
    """A ``.docx``'s text, with Word heading styles emitted as ATX headings."""
    try:
        import docx
    except ImportError:
        return None
    try:
        if path.stat().st_size > _MAX_OFFICE_BYTES:
            return None
        document = docx.Document(str(path))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = _docx_paragraph_text(paragraph)
            if not text:
                continue
            level = _docx_heading_level(getattr(paragraph.style, "name", "") or "")
            parts.append(f"{'#' * level} {text}" if level else text)
        # Tables carry real content in spec documents (API tables, field lists), so keep the
        # cell text as prose rows rather than dropping it.
        for table in document.tables:
            for row in table.rows:
                cells = [" ".join(cell.text.split()) for cell in row.cells]
                line = " · ".join(c for c in cells if c)
                if line:
                    parts.append(line)
    except Exception:  # noqa: BLE001 — encrypted/corrupt docx raises assorted errors
        return None
    return "\n\n".join(parts) or None


def _read_xlsx(path: Path) -> str | None:
    """A workbook's text: one ATX section per sheet, string cells joined as prose.

    Only *string* cells are kept — numbers, dates and formula results are data, not prose
    about code, and would add noise without ever binding. Bounded by sheet and row caps.
    """
    try:
        import openpyxl
    except ImportError:
        return None
    book = None
    try:
        if path.stat().st_size > _MAX_OFFICE_BYTES:
            return None
        book = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        parts: list[str] = []
        for name in book.sheetnames[:_MAX_SHEETS]:
            rows: list[str] = []
            for row in book[name].iter_rows(max_row=_MAX_SHEET_ROWS, values_only=True):
                cells = [" ".join(v.split()) for v in row if isinstance(v, str) and v.strip()]
                if cells:
                    rows.append(" · ".join(cells))
            if rows:
                parts.append(f"# {name}\n" + "\n".join(rows))
    except Exception:  # noqa: BLE001 — encrypted/corrupt workbooks raise assorted errors
        return None
    finally:
        if book is not None:
            book.close()
    return "\n\n".join(parts) or None


# ---- markdown front matter --------------------------------------------------

_FRONT_MATTER_RE = re.compile(r"\A---[ \t]*\r?\n(.*?)\r?\n---[ \t]*(?:\r?\n|\Z)", re.DOTALL)


def _front_matter_prose(block: str) -> str:
    """The *values* of a YAML front-matter block, one per line.

    Keys are dropped deliberately: ``module: billing.invoice`` is a claim about
    ``billing.invoice``, not about the word "module". Keeping keys would feed the binder a
    stream of config-shaped identifiers and inflate drift with things no one wrote as prose.
    Deliberately line-based, not a YAML parse — this needs no dependency, and front matter that
    isn't valid YAML (templated, hand-edited) still yields its values instead of nothing.
    """
    out: list[str] = []
    for line in block.splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#"):
            continue
        if stripped.startswith("- "):
            value = stripped[2:]
        elif ":" in stripped:
            value = stripped.split(":", 1)[1]
        else:
            value = stripped
        value = value.strip().strip("'\"")
        if value:
            out.append(value)
    return "\n".join(out)


def _read_markdown(path: Path) -> str | None:
    """Markdown text, with any YAML front matter reduced to its values.

    Front matter is document *metadata* (``title:``, ``module:``, ``tags:``). Left raw, the
    ``---`` fences and bare ``key:`` names leak into mention extraction as noise; reduced to
    values, a ``module: billing.invoice`` front-matter line binds like the prose it stands for.
    """
    raw = _read_text(path)
    if raw is None:
        return None
    match = _FRONT_MATTER_RE.match(raw)
    if match is None:
        return raw
    prose = _front_matter_prose(match.group(1))
    body = raw[match.end() :]
    return f"{prose}\n\n{body}" if prose else body


# ---- built-in readers -------------------------------------------------------
# Registered here rather than branched in `read_doc_pages`, so a new format is one
# `register_reader` call and touches nothing that already works.

register_reader(DocReader("markdown", frozenset({".md", ".markdown"}), _read_markdown, sections=True))
register_reader(DocReader("text", frozenset({".rst", ".txt"}), _read_text))
register_reader(DocReader("html", frozenset({".html", ".htm"}), _read_html, sections=True))
register_reader(DocReader("pdf", frozenset({".pdf"}), _read_pdf))
register_reader(DocReader("docx", frozenset({".docx"}), _read_docx, sections=True))
register_reader(DocReader("xlsx", frozenset({".xlsx"}), _read_xlsx, sections=True))

# NOT registered: standalone `.yaml`/`.yml`. A repo's YAML is overwhelmingly *configuration*
# (CI workflows, compose files, manifests), not documentation, and ingesting it would corrupt
# the two doc surfaces we just built: `state`'s doc-coverage would climb because a CI file
# happens to name a module, and `doc_drift` would flag every identifier-shaped config value as
# a stale documentation claim. YAML's genuinely documentary case — front matter on a markdown
# doc — is handled by `_read_markdown` above. A repo that wants config ingested can
# `register_reader` a YAML reader itself.


__all__ = ["DocReader", "is_doc_file", "read_doc_pages", "register_reader", "split_sections"]
